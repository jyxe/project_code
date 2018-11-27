from py2neo import Graph,Node,Relationship
import docx
from selenium import webdriver
from win32com import client

#连接neo4j数据库
graph = Graph("bolt://localhost:7687",username="neo4j",password="13820541017")

#创建文档名字
print("请输入pdf文档的名字(以英文命名,不需要后缀)：", end="")
FILE_NAME = input()

#创建word文档
file = docx.Document()
file.add_paragraph("******************************  北京雄恩教育研发中心研发  *******************************"+"\n")
file.add_paragraph("                                                                                        A")

#建立函数从数据数据库中上传数据,当需要同时上传单词数据和文章数据时请使用本函数
# def upload_data_to_database(cypher_word,cypher_text):
#     graph.run(cypher_word)
#     graph.run(cypher_text)

#只上传文章不上传单词时使用本函数
def upload_data_to_database1(cypher_text):
    graph.run(cypher_text)

#只上传单词不上传内容时使用本函数
# def upload_data_to_database2(cypher_word):
#     graph.run(cypher_word)


#建立函数读取数据库中文章的数据
def read_data_from_database(cypher):
    words = []
    data = graph.run(cypher).to_data_frame()
    data = data.values.tolist()

    for i in range (len(data)):
        #从word中添加两段
        file.add_paragraph(data[0][0])
        file.add_paragraph(data[0][1])

        #将文章中的所有大写字母变为小写字母
        string = str(data[0][i]).lower()

        #将文章中的一些标点符号去掉然后进行切分
        string2 = string.replace('.', '', 1000).replace(',', '', 1000).replace('“', '', 1000).replace('”', '',1000).replace('?', '', 1000).replace('’', '', 1000).replace('!', '', 1000).replace('\r\n', '', 1000)
        items = string2.split()

        #去除重复单词
        string3 = " ".join(sorted(set(items)))

        #将去重后的字符串进行切分，默认为按照空格进行切分
        words_different = string3.split()
        for item in words_different:
            words.append(item)

    # print(words)
    return words

#从数据库中读取单词的数据
def read_word_from_database(cypher):
    data_of_word_tolist = []
    data_of_word = graph.run(cypher).to_data_frame()
    data_of_word = data_of_word.values.tolist()

    for i in range (len(data_of_word)):
        data_of_word_tolist.append(data_of_word[i])

    # print(data_of_word_tolist)
    return data_of_word_tolist

#定义函数将文章中的单词与数据库中的数据进行对比，将有的数据列出来写到word当中去
def compare_word_from_database_and_text(data_from_text,data_from_database):
    k=1
    for i in range (len(data_from_text)):
        for j in range (len(data_from_database)):
            if data_from_text[i] == data_from_database[j][0]:
                file.add_paragraph(str(k)+". "+data_from_text[i]+"    "+data_from_database[j][1]+"  "+data_from_database[j][2])
                print(data_from_text[i]+"  "+data_from_database[j][1]+"  "+data_from_database[j][2])
                k+=1

#程序入口
def main():
    #定义接口，使使用者可以很轻松的完成代码的操作
    print("请输入您要上传数据库的文件名称，在上传之前请先确保您已经文件放置在了neo4j的安装目录中的import文件夹下面不要添加后缀")
    print("请现在输入您需要上传的文件名称且确保其名字的正确性：",end="")
    file_name = input()
    print("在请输入您想创建的节点名称(请使用英文命名)：",end="")
    node_name = input()

    print("请输入您想将输出文件定义什么名字(例如：mywork不需要加后缀)：",end="")
    word_file_name = input()

    #定义上传数据的cypher语句
    # cypher_word = ("load csv with headers from 'file:/word2.csv' as row create (n:word1) set n=row")
    cypher_text = ("load csv with headers from 'file:/"+str(file_name)+".csv' as row create (n:"+str(node_name)+") set n=row")

    upload_data_to_database1(cypher_text)

    #开始读取数据库中文章的数据
    cypher = ("match (n:"+node_name+") return n.a_Reading,n.b_Answers")
    data_from_text = read_data_from_database(cypher)
    # print(data_from_text)

    #给word添加一段
    file.add_paragraph("单词注解：")

    #开始读取数据库中单词的数据
    cypher1 = ("match (n:word2) return n.a_word,n.b_alphabet,n.c_meaning")
    data_from_database = read_word_from_database(cypher1)
    # for i in range(len(data_from_word)):
    #     print(data_from_word[i][0])

    #开始执行单词匹配操作
    compare_word_from_database_and_text(data_from_text,data_from_database)

    #将生成的word文档存储在特定的位置
    file.save("D:\\"+word_file_name+".docx")

if __name__ == "__main__":
    main()