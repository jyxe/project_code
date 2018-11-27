from py2neo import Graph,Node,Relationship
import docx
import os
from win32com import client
#连接neo4j数据库
graph = Graph("bolt://localhost:7687",username="neo4j",password="ducunmei2012")

# print("请输入您要上传数据库的文件名称，在上传之前请先确保您已经文件放置在了neo4j的安装目录中的import文件夹下面不要添加后缀")
# print("请现在输入您需要上传的文件名称且确保其名字的正确性：",end="")
# file_name = input()
#
# print("在请输入您想创建的节点名称(请使用英文命名)：",end="")
# node_name = input()
#
# #定义网站驱动
# # browser = webdriver.Chrome()

#创建word文档
file = docx.Document()
file.add_paragraph("******************************  北京雄恩教育研发中心研发  *******************************"+"\n")

#建立函数从数据数据库中上传数据,当需要同时上传单词数据和文章数据时请使用本函数
# def upload_data_to_database(cypher_word,cypher_text):
#     graph.run(cypher_word)
#     graph.run(cypher_text)

#只上传文章不上传单词时使用本函数
def upload_data_to_database1(cypher_text):
    graph.run(cypher_text)

#只上传单词不上传内容时使用本函数
# def upload_data_to_database2(cypher_word):
#     graph.run(cypher_word)


#建立函数读取数据库中文章的数据
def read_data_from_database(cypher):

    #读取数据库的单词
    data_of_word_tolist = []
    data_of_word = graph.run("match (n:English_word) return distinct n.a_word,n.b_alphabet,n.content").to_data_frame()
    data_of_word = data_of_word.values.tolist()

    for i in range(len(data_of_word)):
        data_of_word_tolist.append(data_of_word[i])
    # print(data_of_word_tolist)

    data = graph.run(cypher).to_data_frame()
    data = data.values.tolist()

    for i in range (len(data)):
        words = []
        for j in range (2):
            #从word中添加两段
            file.add_paragraph(data[i][j])
            # file.add_paragraph(data[i][1])

            #将文章中的所有大写字母变为小写字母
            string = str(data[i][j]).lower()

            #将文章中的一些标点符号去掉然后进行切分
            string2 = string.replace('.', '', 1000).replace(',', '', 1000).replace('“', '', 1000).replace('”', '',1000).replace('?', '', 1000).replace('’', '', 1000).replace('!', '', 1000).replace('\r\n', '', 1000)
            items = string2.split()

            #去除重复单词
            string3 = " ".join(sorted(set(items)))

            #将去重后的字符串进行切分，默认为按照空格进行切分
            words_different = string3.split()
            for item in words_different:
                words.append(item)


        items = " ".join(sorted(set(words)))
        item = items.split()
        del item[:12]
        word_different_from_text = []
        for l1 in range(len(item)):
            if (len(item[l1]) > 2 and len(item[l1]) < 18):
                # print(item[i])
                word_different_from_text.append(item[l1])

        file.add_paragraph("单词注解:")
        k=1
        for i in range(len(word_different_from_text)):
            for j in range(len(data_of_word_tolist)):
                if word_different_from_text[i] == data_of_word_tolist[j][0]:
                    file.add_paragraph(str(k)+". "+word_different_from_text[i]+"    "+data_of_word_tolist[j][1]+"  "+data_of_word_tolist[j][2])
                    print(word_different_from_text[i]+"  "+data_of_word_tolist[j][1]+"  "+data_of_word_tolist[j][2])
                    k += 1

# 定义函数，使得生产的word能够转化为PDF格式的文档
def doc_to_pdf(doc_name, pdf_name):

    try:
        word = client.DispatchEx("Word.Application")
        if os.path.exists(pdf_name):
              os.remove(pdf_name)
        worddoc = word.Documents.Open(doc_name, ReadOnly=1)
        worddoc.SaveAs(pdf_name, FileFormat=17)
        worddoc.Close()
        return pdf_name
    except:
        return 1

def main():
    #定义接口，使使用者可以很轻松的完成代码的操作
    # 定义上传数据的cypher语句
    # cypher_word = ("load csv with headers from 'file:/word2.csv' as row create (n:word1) set n=row")
    # cypher_text = ("load csv with headers from 'file:/"+str(file_name)+".csv' as row create (n:"+str(node_name)+") set n=row")

    # upload_data_to_database1(cypher_text)

    #开始读取数据库中文章的数据
    cypher = ("match (n:E2) return n.a_reading,n.b_answer")
    read_data_from_database(cypher)

    #将生成的word文档存储在特定的位置
    file.save("D:\\python_word"+"\\E2.docx")
    doc_name = "D:\\python_word"+"\\E2.docx"
    ftp_name = "D:\\python_word\\E2.pdf"
    doc_to_pdf(doc_name, ftp_name)
    # browser.get("http://localhost:8081/" + str(file_name) + ".pdf")

if __name__ == "__main__":
    main()